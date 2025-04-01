import os
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
import logging
import sys

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def markdown_to_docx(markdown_text, output_filename):
    try:
        doc = Document()
        lines = markdown_text.splitlines()

        i = 0
        while i < len(lines):
            line = lines[i]

            if line.startswith('# '):
                heading = doc.add_heading(line[2:], level=1)
                set_heading_style(heading)
                i += 1
                if i < len(lines) and not lines[i].startswith('#'):
                    add_paragraph_with_bold(doc, lines[i])
                    i += 1
            elif line.startswith('## '):
                heading = doc.add_heading(line[3:], level=2)
                set_heading_style(heading)
                i += 1
                if i < len(lines) and not lines[i].startswith('#'):
                    add_paragraph_with_bold(doc, lines[i])
                    i += 1
            elif line.startswith('### '):
                heading = doc.add_heading(line[4:], level=3)
                set_heading_style(heading)
                i += 1
                if i < len(lines) and not lines[i].startswith('#'):
                    add_paragraph_with_bold(doc, lines[i])
                    i += 1
            elif line.startswith('#### '):
                heading = doc.add_heading(line[5:], level=4)
                set_heading_style(heading)
                i += 1
                if i < len(lines) and not lines[i].startswith('#'):
                    add_paragraph_with_bold(doc, lines[i])
                    i += 1
            elif line.startswith('##### '):
                heading = doc.add_heading(line[6:], level=5)
                set_heading_style(heading)
                i += 1
                if i < len(lines) and not lines[i].startswith('#'):
                    add_paragraph_with_bold(doc, lines[i])
                    i += 1
            elif line.startswith('###### '):
                heading = doc.add_heading(line[7:], level=6)
                set_heading_style(heading)
                i += 1
                if i < len(lines) and not lines[i].startswith('#'):
                    add_paragraph_with_bold(doc, lines[i])
                    i += 1
            else:
                add_paragraph_with_bold(doc, line)
                i += 1

        doc.save(output_filename)
        logging.info(f"Successfully saved document to {output_filename}")
        return True
    except Exception as e:
        logging.error(f"Error in markdown_to_docx: {str(e)}")
        raise

def set_heading_style(heading):
    try:
        for run in heading.runs:
            run.font.name = '仿宋'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.italic = False
    except Exception as e:
        logging.error(f"Error in set_heading_style: {str(e)}")
        raise

def add_paragraph_with_bold(doc, line):
    try:
        paragraph = doc.add_paragraph()
        parts = line.split('**')
        for i, part in enumerate(parts):
            if i % 2 == 1:  # This part is bold
                run = paragraph.add_run(part)
                run.bold = True
                run.font.name = '仿宋'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            else:  # This part is not bold
                run = paragraph.add_run(part)
                run.font.name = '仿宋'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    except Exception as e:
        logging.error(f"Error in add_paragraph_with_bold: {str(e)}")
        raise

def main(company_name, user_demand):
    try:
        # Check if output.txt exists
        if not os.path.exists("output.txt"):
            raise Exception("Please generate document first (output.txt not found)")

        # Read the content
        with open("output.txt", "r", encoding="utf-8") as file:
            content = file.read()

        if not content:
            raise Exception("output.txt is empty")

        # Convert to docx
        output_file = f"output.docx"
        markdown_to_docx(content, output_file)
        
        return f"Successfully converted case study to Word document: {output_file}"

    except Exception as e:
        error_msg = f"Error in document conversion: {str(e)}"
        logging.error(error_msg)
        raise Exception(error_msg)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        logging.error("Usage: python main4.py <company_name> [user_demand]")
        sys.exit(1)
    
    company_name = sys.argv[1]
    user_demand = sys.argv[2] if len(sys.argv) > 2 else ""
    
    try:
        result = main(company_name, user_demand)
        print(result)
    except Exception as e:
        print(f"Error: {str(e)}")
        sys.exit(1)